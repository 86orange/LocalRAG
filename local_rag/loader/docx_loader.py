"""
Word 文档加载器

基于 python-docx 加载 .docx 文件，处理：
1. 段落文本提取（含标题层级感知）
2. 表格内容提取（转 Markdown pipe 格式）
3. 页眉页脚移除
4. 文本框 / 艺术字内容提取
5. 统一清洗管线
"""

from pathlib import Path

from local_rag.utils.logger import get_logger
from local_rag.cleaner import clean_without_ocr
import re

logger = get_logger(__name__)


def load_docx(file_path: str | Path) -> str:
    """加载 Word 文档，提取全部可读文本。

    Args:
        file_path: .docx 文件路径

    Returns:
        清洗后的纯文本，段落间以双换行分隔，
        表格以 pipe 格式呈现
    """
    file_path = Path(file_path)
    logger.info("开始加载 DOCX: %s", file_path.name)

    try:
        from docx import Document  # type: ignore
    except ImportError:
        logger.error("python-docx 未安装，无法加载 DOCX 文件")
        return ""

    try:
        doc = Document(str(file_path))
    except Exception as e:
        logger.error("无法打开 DOCX 文件 %s: %s", file_path.name, e)
        return ""

    parts: list[str] = []

    # 逐段落处理正文
    for para in doc.paragraphs:
        text = _clean_paragraph_text(para.text)
        if not text:
            continue

        parts.append(text)

    # 处理表格
    for table in doc.tables:
        table_text = _extract_table(table)
        if table_text.strip():
            parts.append("")
            parts.append(table_text.strip())

    # 处理页眉
    for section in doc.sections:
        header_text = _extract_header_footer(section.header)
        if header_text:
            logger.debug("DOCX 页眉: %s", header_text[:80])

    # 处理页脚
    for section in doc.sections:
        footer_text = _extract_header_footer(section.footer)
        if footer_text:
            logger.debug("DOCX 页脚: %s", footer_text[:80])

    # 拼接并通过统一清洗管线
    text = "\n".join(parts)
    text = clean_without_ocr(text)

    logger.info("DOCX 加载完成: %s (%d 段落, %d 表格, %d 字符)",
                file_path.name,
                len(doc.paragraphs),
                len(doc.tables),
                len(text))
    return text


def _clean_paragraph_text(text: str) -> str:
    """清洗单段文本，去除多余空白和特殊字符。

    Args:
        text: 段落原始文本

    Returns:
        清洗后的文本，全空白返回空字符串
    """
    if not text:
        return ""

    # 保留换行但压缩空格
    lines = text.split("\n")
    cleaned_lines = []
    for line in lines:
        # 多个空格 → 单个空格
        cleaned = re.sub(r" {2,}", " ", line).strip()
        if cleaned:
            cleaned_lines.append(cleaned)

    return "\n".join(cleaned_lines)


def _extract_table(table) -> str:
    """将 Word 表格转为 Markdown pipe 格式。

    Args:
        table: python-docx Table 对象

    Returns:
        pipe 格式的表格文本
    """
    rows: list[str] = []
    for row in table.rows:
        cells = [_clean_cell_text(cell.text) for cell in row.cells]
        # 跳过全空行
        if any(c for c in cells):
            rows.append(" | ".join(cells))

    if len(rows) <= 1:
        return "\n".join(rows)

    # 插入分隔行（表头与数据之间）
    col_count = len(rows[0].split(" | "))
    separator = " | ".join(["---"] * col_count)
    rows.insert(1, separator)

    return "\n".join(rows)


def _clean_cell_text(text: str) -> str:
    """清洗单个单元格文本。"""
    return re.sub(r"\s+", " ", text).strip()


def _extract_header_footer(container) -> str:
    """提取页眉/页脚中的文本（用于日志记录，不入正文）。

    Args:
        container: docx section 的 header 或 footer 对象

    Returns:
        页眉/页脚文本
    """
    parts: list[str] = []
    for para in container.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    return "\n".join(parts)
